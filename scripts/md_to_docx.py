# -*- coding: utf-8 -*-
"""将论文 Markdown 按学校规范写入 DOCX。

默认行为：
1) 读取仓库根目录下 `毕业论文正文.md`
2) 读取模板 `毕业论文正文.docx`（仅复用页眉页脚/节信息）
3) 生成 `毕业论文正文_排版完成.docx`，不覆盖模板
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt, RGBColor

BLACK = RGBColor(0x00, 0x00, 0x00)

SIZE_3 = Pt(16)   # 三号
SIZE_4 = Pt(14)   # 四号
SIZE_X4 = Pt(12)  # 小四
SIZE_5 = Pt(10.5) # 五号


def apply_page_layout(doc: Document) -> None:
    """学校规范：A4；上3下2左3右2；页眉2；页脚1（cm）。"""
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(2.0)
    section.header_distance = Cm(2.0)
    section.footer_distance = Cm(1.0)


def _clear_paragraph_runs(p) -> None:
    for run in list(p.runs):
        p._p.remove(run._r)


def _append_page_field(p) -> None:
    r = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    r._r.append(begin)
    r._r.append(instr)
    r._r.append(separate)
    r._r.append(end)
    set_run_font(r, east_asia="Times New Roman", ascii_font="Times New Roman", size=SIZE_5)


def apply_header_footer(doc: Document) -> None:
    section = doc.sections[0]

    hp = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
    _clear_paragraph_runs(hp)
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hrun = hp.add_run("杭州电子科技大学本科毕业论文")
    set_run_font(hrun, east_asia="宋体", ascii_font="Times New Roman", size=SIZE_5)

    fp = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
    _clear_paragraph_runs(fp)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _append_page_field(fp)


def clear_document_body(doc: Document) -> None:
    """清空正文，保留 sectPr 以复用模板的页眉页脚引用。"""
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_run_font(
    run,
    *,
    east_asia: str = "宋体",
    ascii_font: str = "Times New Roman",
    size=SIZE_X4,
    bold: bool | None = None,
    underline: bool | None = None,
) -> None:
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = size
    run.font.color.rgb = BLACK
    if bold is not None:
        run.bold = bold
    if underline is not None:
        run.underline = underline


def format_body_paragraph(p, *, first_indent: bool = True, align=WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(20)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Pt(24) if first_indent else Pt(0)
    p.alignment = align


def format_heading_paragraph(p, *, level: int, title_text: str) -> None:
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = 1.0
    pf.space_before = Pt(24)
    pf.space_after = Pt(24)
    pf.first_line_indent = Pt(0)

    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = p.runs[0] if p.runs else p.add_run(title_text)

    if level == 1 and re.fullmatch(r"ABSTRACT", title_text, re.I):
        set_run_font(run, east_asia="Times New Roman", ascii_font="Times New Roman", size=SIZE_3, bold=True)
        return

    if level == 1:
        set_run_font(run, east_asia="黑体", ascii_font="SimHei", size=SIZE_3, bold=False)
    elif level == 2:
        set_run_font(run, east_asia="黑体", ascii_font="SimHei", size=SIZE_4, bold=False)
    else:
        set_run_font(run, east_asia="黑体", ascii_font="SimHei", size=SIZE_X4, bold=False)


def format_caption_paragraph(p) -> None:
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(20)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def set_cell_shading(cell, fill: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_inline_runs(
    p,
    text: str,
    *,
    default_east_asia: str = "宋体",
    default_ascii: str = "Times New Roman",
    default_size=SIZE_X4,
) -> None:
    """支持 **粗体**、`行内代码`、[文本](链接)。"""
    if not text.strip():
        return

    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))", text)
    for part in parts:
        if not part:
            continue

        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            set_run_font(
                run,
                east_asia=default_east_asia,
                ascii_font=default_ascii,
                size=default_size,
                bold=True,
            )
            continue

        if part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            set_run_font(run, east_asia="Consolas", ascii_font="Consolas", size=SIZE_5)
            continue

        if part.startswith("[") and "](" in part and part.endswith(")"):
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", part)
            if m:
                label, url = m.group(1), m.group(2)
                shown = url if label == url else f"{label} ({url})"
                run = p.add_run(shown)
                set_run_font(
                    run,
                    east_asia=default_east_asia,
                    ascii_font=default_ascii,
                    size=default_size,
                    underline=True,
                )
                continue

        run = p.add_run(part)
        set_run_font(run, east_asia=default_east_asia, ascii_font=default_ascii, size=default_size)


def add_body_paragraph(
    doc: Document,
    text: str,
    *,
    first_indent: bool = True,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    east_asia: str = "宋体",
    ascii_font: str = "Times New Roman",
    size=SIZE_X4,
) -> None:
    p = doc.add_paragraph()
    format_body_paragraph(p, first_indent=first_indent, align=align)
    add_inline_runs(
        p,
        text,
        default_east_asia=east_asia,
        default_ascii=ascii_font,
        default_size=size,
    )


def add_cn_keyword_line(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    format_body_paragraph(p, first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    m = re.match(r"^\s*关键词\s*[：:]\s*(.*)$", text)
    rest = m.group(1).strip() if m else text.strip()
    r1 = p.add_run("关键词：")
    set_run_font(r1, east_asia="黑体", ascii_font="SimHei", size=SIZE_X4)
    r2 = p.add_run(rest)
    set_run_font(r2, east_asia="宋体", ascii_font="Times New Roman", size=SIZE_X4)


def add_en_keyword_line(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    format_body_paragraph(p, first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    m = re.match(r"^\s*(Key\s*words|Keywords)\s*:\s*(.*)$", text, re.I)
    label = m.group(1) if m else "Key words"
    rest = m.group(2).strip() if m else text.strip()
    r1 = p.add_run(f"{label}:")
    set_run_font(r1, east_asia="Times New Roman", ascii_font="Times New Roman", size=SIZE_X4, bold=True)
    r2 = p.add_run(rest)
    set_run_font(r2, east_asia="Times New Roman", ascii_font="Times New Roman", size=SIZE_X4)


def is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.count("|") >= 2


def parse_table_row(line: str) -> list[str]:
    cells = [c.strip() for c in line.strip().split("|")]
    if cells and cells[0] == "":
        cells = cells[1:]
    if cells and cells[-1] == "":
        cells = cells[:-1]
    return cells


def is_separator_row(cells: list[str]) -> bool:
    if not cells:
        return False
    return all(re.match(r"^:?-{3,}:?$", c.strip()) for c in cells if c.strip()) or all(
        set(c.replace("-", "").replace(":", "").strip()) <= {""} for c in cells
    )


def is_image_line(line: str) -> re.Match[str] | None:
    return re.match(r"^\s*!\[(.*?)\]\((.*?)\)\s*$", line)


def is_unordered_item(line: str) -> re.Match[str] | None:
    return re.match(r"^\s*[-*]\s+(.*)$", line)


def is_ordered_item(line: str) -> re.Match[str] | None:
    return re.match(r"^\s*\d+\.\s+(.*)$", line)


def is_block_boundary(line: str) -> bool:
    s = line.rstrip()
    if not s.strip():
        return True
    if s.strip() == "---":
        return True
    if s.strip().startswith("```"):
        return True
    if re.match(r"^(#{1,6})\s+", s):
        return True
    if s.startswith("> "):
        return True
    if is_table_row(s):
        return True
    if is_image_line(s):
        return True
    if is_unordered_item(s) or is_ordered_item(s):
        return True
    return False


def is_table_caption_line(line: str) -> bool:
    s = line.strip()
    return bool(re.match(r"^表\s*\d", s) or re.match(r"^表\s*[一二三四五六七八九十]", s))


def add_heading(doc: Document, text: str, *, level: int) -> None:
    p = doc.add_paragraph(text)
    format_heading_paragraph(p, level=level, title_text=text)


def main() -> int:
    root = Path(__file__).resolve().parents[1]
    md_path = root / "毕业论文正文.md"
    out_path = root / "毕业论文正文_排版完成.docx"
    template_path = root / "毕业论文正文.docx"

    if len(sys.argv) >= 2:
        md_path = Path(sys.argv[1])
    if len(sys.argv) >= 3:
        out_path = Path(sys.argv[2])
    if len(sys.argv) >= 4:
        template_path = Path(sys.argv[3])

    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document(str(template_path))

    apply_page_layout(doc)
    apply_header_footer(doc)

    clear_document_body(doc)
    apply_page_layout(doc)
    apply_header_footer(doc)
    section = doc.sections[0]
    max_width = section.page_width - section.left_margin - section.right_margin

    first_hash_title_done = False
    in_code = False
    code_buf: list[str] = []
    section_mode = "body"
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]
        raw = line.rstrip()

        if in_code:
            if raw.strip().startswith("```"):
                p = doc.add_paragraph()
                format_body_paragraph(p, first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
                p.paragraph_format.left_indent = Cm(0.6)
                run = p.add_run("\n".join(code_buf))
                set_run_font(run, east_asia="Consolas", ascii_font="Consolas", size=SIZE_5)
                code_buf = []
                in_code = False
            else:
                code_buf.append(line)
            i += 1
            continue

        if raw.strip().startswith("```"):
            in_code = True
            i += 1
            continue

        if raw.strip() == "---":
            i += 1
            continue

        img_match = is_image_line(raw)
        if img_match:
            alt_text = img_match.group(1).strip()
            rel_path = img_match.group(2).strip()
            img_path = (md_path.parent / rel_path).resolve()

            if img_path.exists():
                pic_para = doc.add_paragraph()
                format_body_paragraph(pic_para, first_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
                run = pic_para.add_run()
                run.add_picture(str(img_path), width=Emu(int(max_width * 0.9)))

                if alt_text:
                    cap_para = doc.add_paragraph()
                    format_caption_paragraph(cap_para)
                    add_inline_runs(cap_para, alt_text, default_size=SIZE_5)
                    for r in cap_para.runs:
                        set_run_font(r, east_asia="宋体", ascii_font="Times New Roman", size=SIZE_5)
            else:
                add_body_paragraph(doc, f"[图片缺失] {rel_path}")

            i += 1
            continue

        if raw.startswith("> "):
            p = doc.add_paragraph()
            format_body_paragraph(p, first_indent=True)
            add_inline_runs(p, raw[2:].strip(), default_size=SIZE_X4)
            i += 1
            continue

        unordered = is_unordered_item(raw)
        if unordered:
            p = doc.add_paragraph(style="List Bullet")
            format_body_paragraph(p, first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
            add_inline_runs(p, unordered.group(1).strip(), default_size=SIZE_X4)
            i += 1
            continue

        ordered = is_ordered_item(raw)
        if ordered:
            p = doc.add_paragraph(style="List Number")
            format_body_paragraph(p, first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
            add_inline_runs(p, ordered.group(1).strip(), default_size=SIZE_X4)
            i += 1
            continue

        if is_table_caption_line(raw) and i + 1 < n and is_table_row(lines[i + 1].strip()):
            cp = doc.add_paragraph()
            format_caption_paragraph(cp)
            run = cp.add_run(raw.strip())
            set_run_font(run, east_asia="宋体", ascii_font="Times New Roman", size=SIZE_5)
            i += 1
            continue

        if is_table_row(raw):
            rows: list[list[str]] = []
            while i < n and is_table_row(lines[i].strip()):
                row = parse_table_row(lines[i])
                if not is_separator_row(row):
                    rows.append(row)
                i += 1

            if not rows:
                continue

            ncol = max(len(r) for r in rows)
            for r in rows:
                while len(r) < ncol:
                    r.append("")

            tbl = doc.add_table(rows=len(rows), cols=ncol)
            tbl.style = "Table Grid"

            for ri, row in enumerate(rows):
                for ci, cell_text in enumerate(row):
                    cell = tbl.rows[ri].cells[ci]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ri == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    add_inline_runs(p, cell_text, default_size=SIZE_5)
                    for r in p.runs:
                        set_run_font(r, size=SIZE_5)
                    if ri == 0:
                        set_cell_shading(cell, "E7E6E6")
            continue

        m1 = re.match(r"^(#{1,6})\s+(.*)$", raw)
        if m1:
            hashes, title = m1.group(1), m1.group(2).strip()
            level = len(hashes)
            is_special_h1 = bool(re.match(r"^(摘\s*要|ABSTRACT|致\s*谢|参考文献|附\s*录|目录)$", title, re.I))

            if level == 1 and not first_hash_title_done:
                # 题目：黑体三号，居中。
                p = doc.add_paragraph(title)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(24)
                p.paragraph_format.first_line_indent = Pt(0)
                run = p.runs[0]
                set_run_font(run, east_asia="黑体", ascii_font="SimHei", size=SIZE_3)
                first_hash_title_done = True
                section_mode = "title"
            else:
                if level == 1 or is_special_h1:
                    pb = doc.add_paragraph()
                    pb.add_run().add_break(WD_BREAK.PAGE)
                    add_heading(doc, title, level=1)
                elif level == 2:
                    add_heading(doc, title, level=2)
                elif level == 3:
                    add_heading(doc, title, level=3)
                else:
                    add_heading(doc, title, level=3)

                if re.match(r"^摘\s*要$", title):
                    section_mode = "cn_abstract"
                elif re.match(r"^ABSTRACT$", title, re.I):
                    section_mode = "en_abstract"
                elif re.match(r"^参考文献$", title):
                    section_mode = "references"
                elif re.match(r"^目录$", title):
                    section_mode = "toc"
                else:
                    section_mode = "body"
            i += 1
            continue

        if not raw.strip():
            i += 1
            continue

        buf = [raw]
        i += 1
        while i < n:
            nxt = lines[i].rstrip()
            if is_block_boundary(nxt):
                break
            buf.append(nxt)
            i += 1

        para_text = " ".join(x.strip() for x in buf)

        if section_mode == "cn_abstract" and re.match(r"^\s*关键词\s*[：:]", para_text):
            add_cn_keyword_line(doc, para_text)
            continue

        if section_mode == "en_abstract" and re.match(r"^\s*(Key\s*words|Keywords)\s*:", para_text, re.I):
            add_en_keyword_line(doc, para_text)
            continue

        if section_mode == "en_abstract":
            add_body_paragraph(
                doc,
                para_text,
                first_indent=True,
                east_asia="Times New Roman",
                ascii_font="Times New Roman",
                size=SIZE_X4,
            )
            continue

        if section_mode in {"references", "toc"} or re.match(r"^\s*\[\d+\]", para_text):
            add_body_paragraph(
                doc,
                para_text,
                first_indent=False,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                east_asia="宋体",
                ascii_font="Times New Roman",
                size=SIZE_X4,
            )
            continue

        add_body_paragraph(doc, para_text, first_indent=True)

    doc.save(out_path)
    print(f"Written: {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
